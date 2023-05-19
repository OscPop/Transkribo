import os
import sys
import argparse
import docx
import pandas as pd
import numpy as np


def read_from_docx(folder_path, file_name):
    """
    Läser in från .docx och gör en .txt-fil av innehållet.
    Kräver att .docx-filen har rätt format.
    """

    if args.verbosity == 1:
        print(f"Reading from .docx ...")

    doc = docx.Document(os.path.join(folder_path, file_name))
    lines = []
    for docpara in doc.paragraphs:
        for line in docpara.text.split("\n"):
            lines.append(line)

    file_name = "teams_transkribering_raw.txt"
    file_path = os.path.join(folder_path, file_name)
    if args.verbosity == 1:
        print(f"Saving file: '{file_name}'")

    with open(file_path, "w", encoding="utf-8") as f:
        for line in lines:
            f.write(line+"\n")



def fix_format(folder_path):
    # Städar upp i transkriberingen
    #   - Skickar ut filer i olika format, beroende på vad man vill ha
    #   - Bygger på att filen följer formatet: 
    #       1. Tid
    #       2. Talare
    #       3. Information


    if args.verbosity == 1:
        print(f"Fixing formatting ...\n")


    input_file = "teams_transkribering_raw.txt"

    with open(os.path.join(folder_path, input_file), "r", encoding="utf-8") as f:
        lines = [line.rstrip() for line in f]

    # Listor för att spara olika information - inte mest minneseffektiva kanske
    lines_grupperad = []
    lines_talare_tid = []
    lines_no_timestamps = []
    lines_no_timestamps_no_names = []
    lines_tsv_grupperad = []
    lines_tsv_ej_grupperad = []
    unique_speakers = []

    # Hålla koll på vem som pratar och vem som pratade sist
    curr_speaker = ""
    last_speaker = ""
    last_milliseconds = 0
    acc_speaker_count = 0

    # Gör en iterator för att använda 'next()'
    iterator = iter(lines)
    for it in iterator:
        
        # Itererar 3 rader i taget; tid, namn, mening
        timestamp, curr_speaker, line = it, next(iterator), next(iterator)
        if curr_speaker not in unique_speakers:
            unique_speakers.append(curr_speaker)

        # Gör om tidsstämplar till millisekunder och räkna med dem istället
        hour, minute, second_millisecond = timestamp.split(" --> ")[0].split(":")
        second, millisecond = second_millisecond.split(".")
        milliseconds1 = int(hour) * 3600000 + int(minute) * 60000 + int(second) * 1000 + int(millisecond)
        hour, minute, second_millisecond = timestamp.split(" --> ")[1].split(":")
        second, millisecond = second_millisecond.split(".")
        milliseconds2 = int(hour) * 3600000 + int(minute) * 60000 + int(second) * 1000 + int(millisecond)
        

        # Om ny talare, skriv ut information om talare och tidsstämpel
        if curr_speaker != last_speaker:
            lines_grupperad.append(curr_speaker)
            lines_talare_tid.append(timestamp)
            lines_talare_tid.append(curr_speaker)
            lines_no_timestamps.append(curr_speaker)
            lines_tsv_grupperad.append([milliseconds1, milliseconds2, curr_speaker, ""])

            lines_tsv_grupperad[acc_speaker_count-1][1] = last_milliseconds
            acc_speaker_count += 1
                

        lines_grupperad.append(line)
        lines_no_timestamps.append(line)
        lines_no_timestamps_no_names.append(line)
        lines_tsv_grupperad[-1][-1] += f"{line} "
        lines_tsv_ej_grupperad.append([milliseconds1, milliseconds2, curr_speaker, line])

        # Sätter förra talaren till nuvarande talare för nästa iteration
        last_speaker = curr_speaker
        last_milliseconds = milliseconds2


    # Skriv texterna till nya filer

    # Talare grupperad - .txt
    with open(os.path.join(folder_path, "teams_grupperad.txt"), "w", encoding="utf-8") as f:
        f.write("\n".join(lines_grupperad))

    # No timestamps
    with open(os.path.join(folder_path, "teams_no_timestamps.txt"), "w", encoding="utf-8") as f:
        f.write("\n".join(lines_no_timestamps))

    # No timestamps no names
    with open(os.path.join(folder_path, "teams_no_timestamps_no_names.txt"), "w", encoding="utf-8") as f:
        f.write("\n".join(lines_no_timestamps_no_names))

    # Talare och tid
    with open(os.path.join(folder_path, "teams_talare_tid.txt"), "w", encoding="utf-8") as f:
        f.write("\n".join(lines_talare_tid))

    # Unika talare
    with open(os.path.join(folder_path, "teams_talare.txt"), "w", encoding="utf-8") as f:
        f.write("\n".join(unique_speakers))


    # Gör en dataframe/.tsv-fil för varje gång någon ny pratar
    df_tsv_grupperad = pd.DataFrame(lines_tsv_grupperad, columns=["start", "end", "speaker", "text"])
    df_tsv_grupperad.to_csv(os.path.join(folder_path, "teams_grupperad.tsv"), sep="\t", index=False)
    #print(f"Shape of df: {df_tsv_grupperad.shape}")

    # En dataframe som inte grupperar för varje ny talare
    df_tsv_ej_grupperad = pd.DataFrame(lines_tsv_ej_grupperad, columns=["start", "end", "speaker", "text"])
    df_tsv_ej_grupperad.to_csv(os.path.join(folder_path, "teams_ej_grupperad.tsv"), sep="\t", index=False, encoding="utf-8")
    #print(f"Shape of df2: {df_tsv_ej_grupperad.shape}")



def write_to_docx(folder_path):
    """
    Skriv den omformatterade texten till .docx
    """
    if args.verbosity == 1:
        print("Writing to .docx ...\n")
    
    with open(os.path.join(folder_path, "teams_grupperad.txt"), "r", encoding="utf-8") as f:
        text = [line.strip() for line in f.readlines()]

    with open(os.path.join(folder_path, "teams_talare.txt"), "r", encoding="utf-8") as f:
        talare = [line.strip() for line in f.readlines()]


    document = docx.Document()

    document.add_heading(f"{folder_path.split('/')[-1]} \nTranskribering - teams", 0)
    p = document.add_paragraph()

    if args.wordcloud:
        document.add_picture(os.path.join(folder_path, "wordcloud.png"), width=docx.shared.Inches(6))
        document.add_page_break()


    for line in text:
        if line in talare:
            run = p.add_run()
            run.add_break()
            run.add_break()
            p.add_run(str(line)).bold = True
            run = p.add_run()
            run.add_break()
        else:
            run = p.add_run(str(line)+" ")
            #font = run.font
            #font.color.rgb = shared.RGBColor(200, 0, 0)

    file_name = f"{folder_path.split('/')[-1]} Transkribering - teams.docx"

    if args.verbosity == 1:
        print(f"Saving file '{file_name}' to folder '{folder_path}'")

    file_path = os.path.join(folder_path, file_name)
    document.save(file_path)


def send_to_zip():
    pass


class NoDocxFileFound(Exception):
    """Ingen .docx fil hittades i cwd."""


if __name__ == "__main__":
    print("\n\nRunning program ...\n")
    parser = argparse.ArgumentParser(description="Formatterar om teams transkriberingar från .docx till .docx.", 
                                     epilog="Use quotes surrounding the directory name.\n",
                                     formatter_class=argparse.ArgumentDefaultsHelpFormatter,
                                     add_help=True)
    parser.add_argument("-d", "--dir",type=str, help="Which directory to work with.", required=True)
    parser.add_argument("-v", "--verbosity", type=int, help="Control the verbosity of the output", choices=[0,1], default=1)
    parser.add_argument("-w", "--wordcloud", type=bool, help="Inkludera wordcloud i dokumentet", default=False, choices=[True,False])

    args = parser.parse_args()

    # Ändra cwd till specificerad dir
    folder_path = args.dir

    # Leta efter .docx-fil
    input_docx_file = None
    num_docx = 0
    for file in os.listdir(folder_path):
        if file.endswith(".docx"):
            input_docx_file = file
            num_docx += 1
            print(f"Found .docx file: '{input_docx_file}'")

    # Om inte hittar .docx - avbryt
    if not input_docx_file:
        print("No .docx file found in specified directory ...")
        raise NoDocxFileFound
        #sys.exit(1)
    elif num_docx > 1:
        print("Multiple .docx files found ...")
        print("Defaulting to standard file if present ...")
        input_docx_file = "teams_transkribering_raw.docx"


    read_from_docx(folder_path, input_docx_file)

    fix_format(folder_path)

    write_to_docx(folder_path)

